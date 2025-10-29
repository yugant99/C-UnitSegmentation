#!/usr/bin/env python3
from __future__ import annotations

import argparse
import sys
from pathlib import Path

try:
    from docx import Document  # python-docx
except Exception as exc:  # pragma: no cover
    print(
        "Missing dependency 'python-docx'. Install with: pip install python-docx",
        file=sys.stderr,
    )
    raise


def extract_docx_text(docx_path: Path) -> str:
    """Extract text (paragraphs + tables) from a .docx file.

    Paragraphs are joined by newlines. Tables (if any) are appended with rows
    rendered as tab-separated cells, also newline-delimited.
    """
    document = Document(str(docx_path))

    lines: list[str] = []

    # Paragraph text
    for paragraph in document.paragraphs:
        lines.append(paragraph.text)

    # Table text (if present)
    for table in document.tables:
        for row in table.rows:
            cell_texts = [cell.text for cell in row.cells]
            lines.append("\t".join(cell_texts))

    # Ensure file ends with a single trailing newline
    text = "\n".join(lines).rstrip("\n") + "\n"
    return text


def map_output_name(stem: str) -> str:
    """Map a DOCX filename stem to the normalized TXT filename expected by the pipeline.

    - Raw variants (D/before Segmentation) → (Descript generated).txt
    - Gold variants (Final/after Segmentation) → (Orthographic Segmented Transcript).txt
    - Otherwise: keep stem + .txt
    """
    name = stem.strip()

    # Raw variants → (Descript generated)
    if name.endswith(" (D)"):
        return name[:-4] + " (Descript generated).txt"
    if name.endswith(" - before Segmentation"):
        return name[:-22] + " (Descript generated).txt"

    # Gold variants → (Orthographic Segmented Transcript)
    if name.endswith(" - Final"):
        return name[:-8] + " (Orthographic Segmented Transcript).txt"
    if name.endswith(" - after Segmentation"):
        return name[:-21] + " (Orthographic Segmented Transcript).txt"

    # Fallback: keep stem
    return name + ".txt"


def convert_directory(input_dir: Path, output_dir: Path) -> list[Path]:
    """Convert all .docx files in input_dir to .txt files in output_dir.

    Returns a list of written file paths.
    """
    if not input_dir.exists() or not input_dir.is_dir():
        raise FileNotFoundError(f"Input directory not found: {input_dir}")

    output_dir.mkdir(parents=True, exist_ok=True)

    written: list[Path] = []
    for docx_file in sorted(input_dir.glob("*.docx")):
        try:
            text = extract_docx_text(docx_file)
        except Exception as exc:  # pragma: no cover
            print(f"[WARN] Failed to read '{docx_file.name}': {exc}", file=sys.stderr)
            continue

        out_filename = map_output_name(docx_file.stem)
        out_path = output_dir / out_filename
        try:
            out_path.write_text(text, encoding="utf-8")
            written.append(out_path)
            print(f"[OK] Wrote {out_path}")
        except Exception as exc:  # pragma: no cover
            print(f"[WARN] Failed to write '{out_path.name}': {exc}", file=sys.stderr)

    if not written:
        print("[INFO] No .docx files found to convert.")

    return written


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Extract text from .docx files to .txt")
    parser.add_argument(
        "--input-dir",
        default="/Users/yuganthareshsoni/CunitSegementation/Be EPIC-VR transcript examplesV02",
        help="Absolute path to directory containing .docx files",
    )
    parser.add_argument(
        "--output-dir",
        default="/Users/yuganthareshsoni/CunitSegementation/extracted_text",
        help="Directory to write extracted .txt files",
    )

    args = parser.parse_args(argv)
    input_dir = Path(args.input_dir).expanduser()
    output_dir = Path(args.output_dir).expanduser()

    try:
        written = convert_directory(input_dir, output_dir)
    except Exception as exc:
        print(f"[ERROR] {exc}", file=sys.stderr)
        return 1

    if not written:
        return 2
    return 0


if __name__ == "__main__":  # pragma: no cover
    raise SystemExit(main())


